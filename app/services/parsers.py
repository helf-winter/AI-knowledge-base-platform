from __future__ import annotations

import csv
import io
import re

from fastapi import UploadFile

from app.core.config import get_settings
from app.core.exceptions import ValidationAppError

settings = get_settings()

try:
    import pytesseract
    if settings.tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd
except Exception:
    pytesseract = None

ALLOWED_EXTS = {"pdf", "docx", "png", "jpg", "jpeg", "txt", "md", "csv", "xlsx", "xls"}


def sanitize_text_for_storage(text: str) -> str:
    """Remove control characters that PostgreSQL text fields cannot store."""
    if not text:
        return ""
    text = text.replace("\x00", "")
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f]", " ", text)
    return text.strip()


def validate_upload_file(file: UploadFile) -> None:
    if not file.filename:
        raise ValidationAppError("文件名不能为空")
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTS:
        raise ValidationAppError(f"不支持的文件类型: {ext}")


def _extract_text_from_docx(content: bytes) -> str:
    try:
        from docx import Document
    except Exception:
        return "DOCX 解析依赖未安装，已接收文件。"

    try:
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        return sanitize_text_for_storage("\n".join(parts)) or "DOCX 文件已接收，但未提取出文本。"
    except Exception as exc:
        return f"DOCX 解析失败：{exc}"


def _extract_text_from_csv(content: bytes) -> str:
    try:
        text = sanitize_text_for_storage(content.decode("utf-8", errors="ignore"))
        reader = csv.reader(io.StringIO(text))
        rows = []
        for row in reader:
            if any(cell.strip() for cell in row):
                rows.append(" | ".join(cell.strip() for cell in row))
        return sanitize_text_for_storage("\n".join(rows)) or "CSV 文件已接收，但未提取出内容。"
    except Exception as exc:
        return f"CSV 解析失败：{exc}"


def _extract_text_from_excel(content: bytes) -> str:
    try:
        import pandas as pd
    except Exception:
        return "Excel 解析依赖未安装，已接收文件。"

    try:
        bio = io.BytesIO(content)
        sheets = pd.read_excel(bio, sheet_name=None)
        parts: list[str] = []
        for sheet_name, df in sheets.items():
            parts.append(f"[sheet] {sheet_name}")
            if df.empty:
                parts.append("空表")
                continue
            sample = df.astype(str).fillna("").head(50)
            for _, row in sample.iterrows():
                row_text = " | ".join(str(v).strip() for v in row.tolist() if str(v).strip())
                if row_text:
                    parts.append(row_text)
        return sanitize_text_for_storage("\n".join(parts)) or "Excel 文件已接收，但未提取出内容。"
    except Exception as exc:
        return f"Excel 解析失败：{exc}"


def _extract_text_from_pdf(content: bytes, file_name: str) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        return f"{file_name} 已接收，但当前环境缺少 PDF 文本解析依赖 pypdf。"

    try:
        reader = PdfReader(io.BytesIO(content))
        parts: list[str] = []
        for index, page in enumerate(reader.pages[:200], start=1):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""
            page_text = sanitize_text_for_storage(page_text)
            if page_text:
                parts.append(f"[page {index}]\n{page_text}")
        return sanitize_text_for_storage("\n\n".join(parts)) or f"{file_name} 已接收，但未能从 PDF 中提取到可用文本。"
    except Exception as exc:
        return f"PDF 解析失败：{exc}"


def _extract_text_from_image(content: bytes) -> str:
    try:
        from PIL import Image
    except Exception:
        return "图片 OCR 依赖未安装，已接收图片文件。"

    try:
        image = Image.open(io.BytesIO(content))
        try:
            import pytesseract
            text = pytesseract.image_to_string(image, lang="chi_sim+eng")
            if text.strip():
                return sanitize_text_for_storage(text)
            return "图片已接收，但未识别出可用文字。"
        except Exception as exc:
            return f"图片 OCR 不可用：{exc}。已接收图片文件。"
    except Exception as exc:
        return f"图片解析失败：{exc}"


def extract_text_from_bytes(file_name: str, content: bytes) -> str:
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else "txt"
    if not content:
        raise ValidationAppError("文件内容不能为空")

    if ext in {"txt", "md"}:
        return sanitize_text_for_storage(content.decode("utf-8", errors="ignore"))
    if ext == "pdf":
        return _extract_text_from_pdf(content, file_name)
    if ext == "docx":
        return _extract_text_from_docx(content)
    if ext == "csv":
        return _extract_text_from_csv(content)
    if ext in {"xlsx", "xls"}:
        return _extract_text_from_excel(content)
    if ext in {"png", "jpg", "jpeg"}:
        return _extract_text_from_image(content)

    text = sanitize_text_for_storage(content.decode("utf-8", errors="ignore"))
    if text:
        return text
    return f"{file_name} 的二进制内容已接收，后续可接入专用解析器。"


def chunk_text(text: str, chunk_size_tokens: int, overlap_tokens: int) -> list[str]:
    text = sanitize_text_for_storage(text)
    if not text:
        return []
    words = text.split()
    if len(words) <= chunk_size_tokens:
        return [text]

    chunks: list[str] = []
    start = 0
    step = max(1, chunk_size_tokens - overlap_tokens)
    while start < len(words):
        end = min(len(words), start + chunk_size_tokens)
        chunk_words = words[start:end]
        if chunk_words:
            chunk = sanitize_text_for_storage(" ".join(chunk_words))
            if chunk:
                chunks.append(chunk)
        if end >= len(words):
            break
        start += step
    return chunks
